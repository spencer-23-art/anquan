import json
import re
import uuid
from datetime import datetime
from decimal import Decimal
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image
from sqlalchemy.orm import Session

from app.config import settings
from app.core.uploads import validate_image_content
from app.models.fine_ticket import FineTicketType
from app.services import ai_config_service

DOCS_DIR = Path(settings.UPLOAD_DIR) / "fines" / "docs"
PHOTOS_DIR = Path(settings.UPLOAD_DIR) / "fines" / "photos"
COUNTER_FILE = Path(settings.UPLOAD_DIR) / "fines" / "counter.json"

TITLE_RED = RGBColor(165, 0, 0)
TEXT_BLACK = RGBColor(0, 0, 0)
LABEL_FILL = "F2F2F2"
BOX_RED = "C00000"
BOX_BLACK = "000000"


def ensure_storage_dirs() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    PHOTOS_DIR.mkdir(parents=True, exist_ok=True)
    COUNTER_FILE.parent.mkdir(parents=True, exist_ok=True)


def _load_counter() -> dict[str, int]:
    ensure_storage_dirs()
    if COUNTER_FILE.exists():
        return json.loads(COUNTER_FILE.read_text(encoding="utf-8"))
    return {}


def _save_counter(counter: dict[str, int]) -> None:
    ensure_storage_dirs()
    COUNTER_FILE.write_text(
        json.dumps(counter, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def peek_next_number(ticket_type: FineTicketType) -> str:
    today = datetime.now().strftime("%Y%m%d")
    prefix = "ZL" if ticket_type == FineTicketType.QUALITY else "AQ"
    counter = _load_counter()
    key = f"{prefix}-{today}"
    current = counter.get(key, 0) + 1
    return f"{prefix}-{today[:4]}-{today[4:8]}-{current:03d}"


def consume_next_number(ticket_type: FineTicketType) -> str:
    today = datetime.now().strftime("%Y%m%d")
    prefix = "ZL" if ticket_type == FineTicketType.QUALITY else "AQ"
    counter = _load_counter()
    key = f"{prefix}-{today}"
    current = counter.get(key, 0) + 1
    counter[key] = current
    _save_counter(counter)
    return f"{prefix}-{today[:4]}-{today[4:8]}-{current:03d}"


def amount_to_chinese(amount: Decimal) -> str:
    digits = ["零", "壹", "贰", "叁", "肆", "伍", "陆", "柒", "捌", "玖"]
    small_units = ["", "拾", "佰", "仟"]
    big_units = ["", "万", "亿", "兆"]
    amount = Decimal(amount).quantize(Decimal("0.01"))
    integer_part = int(amount)
    decimal_part = int((amount - integer_part) * 100)

    if integer_part == 0:
        integer_text = "零"
    else:
        integer_text = ""
        groups: list[str] = []
        value = integer_part
        while value > 0:
            groups.append(f"{value % 10000:04d}")
            value //= 10000

        zero_pending = False
        for group_index in range(len(groups) - 1, -1, -1):
            group = groups[group_index]
            group_value = int(group)
            if group_value == 0:
                zero_pending = True
                continue

            group_text = ""
            local_zero = False
            for index, char in enumerate(group):
                digit = int(char)
                unit_index = 3 - index
                if digit == 0:
                    if group_text:
                        local_zero = True
                    continue
                if (zero_pending or local_zero) and group_text:
                    group_text += "零"
                group_text += digits[digit] + small_units[unit_index]
                zero_pending = False
                local_zero = False

            integer_text += group_text + big_units[group_index]
            if group_value < 1000 and group_index > 0:
                zero_pending = True

    if decimal_part == 0:
        return f"{integer_text}元整"

    jiao = decimal_part // 10
    fen = decimal_part % 10
    decimal_text = ""
    if jiao:
        decimal_text += f"{digits[jiao]}角"
    if fen:
        decimal_text += f"{digits[fen]}分"
    return f"{integer_text}元{decimal_text}"


def _ticket_title(ticket_type: FineTicketType) -> str:
    if ticket_type == FineTicketType.QUALITY:
        return "工程质量罚款通知单"
    return "工程安全罚款通知单"


def _ticket_name(ticket_type: FineTicketType) -> str:
    if ticket_type == FineTicketType.QUALITY:
        return "质量"
    return "安全"


def _fallback_description(
    *,
    user_input: str,
    project_name: str,
    team_name: str,
    location: str,
    discovery_date: str,
    ticket_type: FineTicketType,
) -> str:
    ticket_name = _ticket_name(ticket_type)
    rule_reference = _resolve_rule_reference(user_input, ticket_type)
    fact_date = _format_fact_date(discovery_date)
    risk_tail = _resolve_risk_tail(user_input, ticket_type)
    return (
        f"一、违章事实及经过：{fact_date}，经现场{ticket_name}检查，项目“{project_name}”中，"
        f"{team_name}在“{location}”作业期间存在{user_input}等问题，反映出班组现场管理和作业控制落实不到位。"
        f"{risk_tail}现场已明确提出整改要求，并责令相关人员立即整改。\n\n"
        f"二、违反条款及性质：上述行为已违反{rule_reference}及项目现场{ticket_name}管理要求，"
        "属于应当立即整改的违规行为。若不及时纠正，极易导致现场风险进一步扩大，性质较为严重。"
    )


def generate_description(
    db: Session,
    user_input: str,
    project_name: str,
    team_name: str,
    location: str,
    discovery_date: str,
    ticket_type: FineTicketType,
) -> str:
    rule_reference = resolve_rule_reference(
        user_input=user_input,
        project_name=project_name,
        team_name=team_name,
        location=location,
        ticket_type=ticket_type,
    )
    fallback = _fallback_description(
        user_input=user_input,
        project_name=project_name,
        team_name=team_name,
        location=location,
        discovery_date=discovery_date,
        ticket_type=ticket_type,
    )

    ticket_name = _ticket_name(ticket_type)
    fact_date = _format_fact_date(discovery_date)
    system_prompt = (
        f"你是一名建筑工程{ticket_name}管理负责人，擅长起草正式的施工罚款通知单内容。"
        "输出必须是中文正式公文风格，内容严肃、明确、可直接粘贴进罚款通知单。"
    )
    user_prompt = (
        "请根据以下信息生成两段式罚款描述，严格遵守要求：\n"
        "1. 仅输出中文正文，不要使用 markdown。\n"
        "2. 必须保留两个段落标题：\n"
        "一、违章事实及经过：\n"
        "二、违反条款及性质：\n"
        "3. 第一段必须使用提供的发现日期，不得写成X月X日、某日、近日或自行编造其他日期。\n"
        "4. 第一段写清地点、班组、行为经过、现场隐患和后果。\n"
        "5. 第二段必须明确写出违反的规范名称，不得省略规范依据。\n"
        "6. 第二段必须优先引用与本次违规直接相关的专业规范、规程或管理规定；通用规范只能作为兜底补充，不得所有场景都套用同一个GB 55034-2022。\n"
        "7. 不要编造无法确认的具体条号；无法确定条号时，使用“相关条款关于……的规定”表述。\n"
        "8. 两段都要写成完整句，不允许只写标题、日期、短语或空泛套话。\n"
        "9. 行文要自然、正式，避免每次都使用完全相同的套话。\n"
        "10. 第一段和第二段都应尽量写得充实，适合直接用于罚款通知单。\n\n"
        f"项目名称：{project_name}\n"
        f"受罚班组：{team_name}\n"
        f"违规部位：{location}\n"
        f"发现日期：{fact_date}\n"
        f"优先写入的规范依据：{rule_reference}\n"
        f"用户描述：{user_input}\n"
        f"罚单类型：{ticket_name}罚款单"
    )

    # The model is limited to fact wording. The compliance paragraph is deterministic
    # so that it always contains the vetted rule selected from the ticket context.
    system_prompt = "你是施工现场管理人员，只起草罚单中的违章事实。不得编造法规、条款号、整改措施或结束语。"
    user_prompt = (
        "请将以下现场信息写成一段正式、客观的中文违章事实，保留项目、班组、部位和发现日期。"
        "只输出一段正文，不要标题、编号、法规依据、整改建议或任何附加说明。\n"
        f"项目：{project_name}\n班组：{team_name}\n部位：{location}\n发现日期：{fact_date}\n现场描述：{user_input}"
    )

    try:
        _used_provider, payload = ai_config_service.request_chat_completion(
            db,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.1,
            max_tokens=360,
            timeout=60.0,
        )
        content = payload["choices"][0]["message"]["content"].strip()
        if not content:
            return "\n".join(_build_controlled_description(fallback, discovery_date, ticket_type, rule_reference))
        paragraphs = _build_controlled_description(content, discovery_date, ticket_type, rule_reference)
        if _description_too_simple(
            paragraphs,
            project_name=project_name,
            team_name=team_name,
            location=location,
        ):
            paragraphs = _build_controlled_description(fallback, discovery_date, ticket_type, rule_reference)
        return "\n".join(paragraphs)
    except Exception:
        return "\n".join(_build_controlled_description(fallback, discovery_date, ticket_type, rule_reference))


def _set_run_font(
    run,
    *,
    size: int = 12,
    bold: bool = False,
    color: RGBColor | None = None,
    name: str = "Microsoft YaHei",
):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color is not None:
        run.font.color.rgb = color


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def _set_table_grid(table, color: str = BOX_BLACK, size: int = 8) -> None:
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(
                cell,
                top={"val": "single", "sz": size, "color": color},
                left={"val": "single", "sz": size, "color": color},
                bottom={"val": "single", "sz": size, "color": color},
                right={"val": "single", "sz": size, "color": color},
            )


def _set_row_height(row, height_cm: float) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(int(height_cm * 567)))
    tr_pr.append(tr_height)


def _clear_cell(cell) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def _set_cell_text(
    cell,
    text: str,
    *,
    size: int = 12,
    bold: bool = False,
    color: RGBColor = TEXT_BLACK,
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
    fill: str | None = None,
) -> None:
    _clear_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    _set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        _set_cell_shading(cell, fill)


def _add_section_heading(doc: Document, title: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    square = paragraph.add_run("■")
    _set_run_font(square, size=13, bold=True, color=TITLE_RED)
    gap = paragraph.add_run(" ")
    _set_run_font(gap, size=12, color=TITLE_RED)
    text_run = paragraph.add_run(title)
    _set_run_font(text_run, size=15, bold=True, color=TITLE_RED)


def _add_bottom_line(doc: Document) -> None:
    line_table = doc.add_table(rows=1, cols=1)
    line_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    line_table.autofit = False
    line_table.columns[0].width = Cm(16.8)
    cell = line_table.cell(0, 0)
    _clear_cell(cell)
    _set_cell_border(
        cell,
        bottom={"val": "single", "sz": 12, "color": BOX_RED},
    )


def _normalize_description(description: str, ticket_type: FineTicketType) -> list[str]:
    raw = (description or "").strip()
    rule_reference = _default_rule_reference(ticket_type)
    if not raw:
        ticket_name = _ticket_name(ticket_type)
        return [
            f"一、违章事实及经过：经现场{ticket_name}检查，发现存在违规作业行为，已责令立即整改。",
            f"二、违反条款及性质：上述行为违反{rule_reference}及项目{ticket_name}管理要求，存在较大现场风险，性质较为严重。",
        ]

    lines = [line.strip() for line in raw.replace("\r", "").split("\n") if line.strip()]
    if any(line.startswith("一、") or line.startswith("二、") for line in lines):
        return _ensure_clause_paragraph(lines, ticket_type)

    ticket_name = _ticket_name(ticket_type)
    return _ensure_clause_paragraph([
        f"一、违章事实及经过：{raw}",
        f"二、违反条款及性质：上述行为违反{rule_reference}及项目{ticket_name}管理要求，属于应立即整改的违规行为。",
    ], ticket_type)


def _format_display_date(date_str: str | None) -> tuple[str, str]:
    if not date_str:
        current = datetime.now()
        return current.strftime("%Y 年 %m 月 %d 日"), current.strftime("%Y 年 %m 月 %d 日")

    for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"):
        try:
            parsed = datetime.strptime(date_str, fmt)
            display = parsed.strftime("%Y 年 %m 月 %d 日")
            return display, display
        except ValueError:
            continue
    return date_str, date_str


def _format_fact_date(date_str: str | None) -> str:
    display, _ = _format_display_date(date_str)
    return display if display else "当日"


def _current_issue_date() -> str:
    return datetime.now().strftime("%Y 年 %m 月 %d 日")


def _compact_text(text: str, limit: int = 120) -> str:
    cleaned = " ".join(str(text or "").replace("\r", " ").replace("\n", " ").split())
    if len(cleaned) <= limit:
        return cleaned
    return cleaned[: limit - 1].rstrip("，。、；：,. ") + "。"


def _default_rule_reference(ticket_type: FineTicketType) -> str:
    if ticket_type == FineTicketType.QUALITY:
        return "《建筑与市政工程施工质量控制通用规范》GB 55032-2022及项目质量管理制度"
    return "《建筑与市政施工现场安全卫生与职业健康通用规范》GB 55034-2022及项目安全管理制度"


def _has_any(text: str, keywords: Iterable[str]) -> bool:
    return any(keyword in text for keyword in keywords)


RULE_CATALOG: tuple[tuple[FineTicketType, tuple[str, ...], str], ...] = (
    (
        FineTicketType.SAFETY,
        ("安全帽", "未戴帽", "未佩戴帽", "安全鞋", "反光衣", "防护眼镜", "防护手套"),
        "《建筑施工作业劳动防护用品配备及使用标准》JGJ 184-2009中关于个人劳动防护用品配备和正确使用的相关规定",
    ),
    (
        FineTicketType.SAFETY,
        ("脚手架", "脚手板", "连墙件", "扫地杆", "剪刀撑", "架体"),
        "《建筑施工扣件式钢管脚手架安全技术规范》JGJ 130-2011中关于架体基础、连墙件、剪刀撑和作业层防护的相关规定",
    ),
    (
        FineTicketType.SAFETY,
        ("高处", "登高", "安全带", "临边", "洞口", "坠落", "高空"),
        "《建筑施工高处作业安全技术规范》JGJ 80-2016中关于临边洞口防护、登高作业和防坠落措施的相关规定",
    ),
    (
        FineTicketType.SAFETY,
        ("临电", "配电箱", "漏保", "电缆", "接地", "电线", "三级配电", "二级保护"),
        "《施工现场临时用电安全技术规范》JGJ 46-2005中关于配电箱、漏电保护、接地和电缆敷设的相关规定",
    ),
    (
        FineTicketType.SAFETY,
        ("动火", "焊接", "切割", "明火", "气瓶", "乙炔", "氧气瓶"),
        "《建设工程施工现场消防安全技术规范》GB 50720-2011中关于动火审批、消防器材配置和可燃物清理的相关规定",
    ),
    (
        FineTicketType.SAFETY,
        ("吊装", "起重", "信号工", "司索", "塔吊", "吊物", "吊篮"),
        "《建筑施工起重吊装工程安全技术规范》JGJ 276-2012中关于吊装指挥、索具、吊物和作业区域控制的相关规定",
    ),
    (
        FineTicketType.QUALITY,
        ("钢筋", "箍筋", "保护层", "混凝土", "振捣", "蜂窝", "麻面", "露筋"),
        "《混凝土结构工程施工质量验收规范》GB 50204-2015中关于钢筋、混凝土成型质量和检验控制的相关规定",
    ),
    (
        FineTicketType.QUALITY,
        ("钢结构", "钢梁", "钢柱", "焊缝", "高强螺栓", "防腐", "防火涂料"),
        "《钢结构工程施工质量验收标准》GB 50205-2020中关于钢结构安装、连接质量和验收的相关规定",
    ),
    (
        FineTicketType.QUALITY,
        ("抹灰", "空鼓", "开裂", "饰面", "墙砖", "地砖", "平整度"),
        "《建筑装饰装修工程质量验收标准》GB 50210-2018中关于抹灰、饰面和安装工程质量验收的相关规定",
    ),
    (
        FineTicketType.QUALITY,
        ("防水", "渗漏", "漏水", "卷材", "涂膜", "闭水"),
        "《建筑与市政工程防水通用规范》GB 55030-2022中关于防水材料、节点和渗漏控制的相关规定",
    ),
)


def resolve_rule_reference(
    *,
    user_input: str,
    project_name: str = "",
    team_name: str = "",
    location: str = "",
    ticket_type: FineTicketType,
) -> str:
    """Choose a vetted rule from the complete ticket context, never from AI output."""
    context = " ".join([user_input or "", project_name or "", team_name or "", location or ""])
    return _resolve_rule_reference(context, ticket_type)


def _resolve_rule_reference(user_input: str, ticket_type: FineTicketType) -> str:
    base = _default_rule_reference(ticket_type)
    text = str(user_input or "")

    catalog_matches = [
        (sum(keyword in text for keyword in keywords), reference)
        for rule_type, keywords, reference in RULE_CATALOG
        if rule_type == ticket_type and _has_any(text, keywords)
    ]
    if catalog_matches:
        return max(catalog_matches, key=lambda item: item[0])[1]

    if ticket_type == FineTicketType.QUALITY:
        if _has_any(text, ["钢结构", "钢构", "钢梁", "钢柱", "焊缝", "高强螺栓", "防腐", "防火涂料"]):
            return "《钢结构工程施工质量验收标准》GB 50205-2020中关于钢结构安装、连接质量及观感验收的相关条款，及项目质量管理制度"
        if _has_any(text, ["钢筋", "箍筋", "保护层", "混凝土", "振捣", "蜂窝", "麻面", "露筋"]):
            return "《混凝土结构工程施工质量验收规范》GB 50204-2015中关于钢筋、混凝土成型质量和验收控制的相关条款，及项目质量管理制度"
        if _has_any(text, ["砌体", "砌筑", "灰缝", "拉结筋", "构造柱", "砖墙"]):
            return "《砌体结构工程施工质量验收规范》GB 50203-2011中关于砌体施工质量、拉结和灰缝控制的相关条款，及项目质量管理制度"
        if _has_any(text, ["防水", "渗漏", "漏水", "卷材", "涂膜", "闭水"]):
            return "《建筑与市政工程防水通用规范》GB 55030-2022中关于防水工程材料、节点和渗漏控制的相关条款，及项目质量管理制度"
        return base

    if _has_any(text, ["钢结构", "钢构", "立柱", "钢柱", "钢梁", "缆风绳", "缆绳", "临时支撑", "斜撑", "校正"]):
        return "《钢结构工程施工规范》GB 50755-2012中关于钢结构安装临时固定、支撑稳定和缆风绳安全控制的相关条款，及项目安全管理制度"
    if _has_any(text, ["脚手架", "脚手板", "连墙件", "扫地杆", "立杆", "横杆", "剪刀撑", "步距", "跨距", "架体"]):
        return "《建筑施工扣件式钢管脚手架安全技术规范》JGJ 130-2011中关于架体基础、连墙件、剪刀撑和作业层防护的相关条款，及项目安全管理制度"
    if _has_any(text, ["基坑", "地坑", "沟槽", "边坡", "放坡", "支护", "降水", "坍塌"]):
        return "《建筑基坑支护技术规程》JGJ 120-2012中关于基坑支护、边坡稳定、临边防护和监测的相关条款，及项目安全管理制度"
    if _has_any(text, ["模板", "支模", "支撑架", "满堂架", "承重架"]):
        return "《建筑施工模板安全技术规范》JGJ 162-2008中关于模板支撑体系搭设、验收和荷载控制的相关条款，及项目安全管理制度"
    if _has_any(text, ["培训证", "培训证过期", "操作证", "特种作业证", "上岗证", "证书过期", "持证", "证件过期", "无证上岗", "特种作业"]):
        return "《建筑施工特种作业人员管理规定》及《特种作业人员安全技术培训考核管理规定》中关于特种作业人员持证上岗和证件有效性的相关条款，及项目安全管理制度"
    if _has_any(text, ["吊装", "起重", "信号工", "司索", "塔吊", "吊物", "吊篮"]):
        return "《建筑施工起重吊装工程安全技术规范》JGJ 276-2012中关于起重吊装指挥、索具、吊物和作业区域控制的相关条款，及项目安全管理制度"
    if _has_any(text, ["高处", "登高", "安全带", "临边", "洞口", "坠落", "梯子", "高空"]):
        return "《建筑施工高处作业安全技术规范》JGJ 80-2016中关于临边洞口防护、登高作业和防坠落措施的相关条款，及项目安全管理制度"
    if _has_any(text, ["临电", "配电箱", "漏保", "电缆", "接地", "电线", "三级配电", "二级保护"]):
        return "《施工现场临时用电安全技术规范》JGJ 46-2005中关于配电箱、漏电保护、接地接零和电缆敷设的相关条款，及项目安全管理制度"
    if _has_any(text, ["动火", "焊接", "切割", "明火", "气瓶", "乙炔", "氧气瓶"]):
        return "《建设工程施工现场消防安全技术规范》GB 50720-2011中关于动火审批、消防器材配置和可燃物清理的相关条款，及项目安全管理制度"
    if _has_any(text, ["安全帽", "防护用品", "劳保", "反光衣", "防护眼镜", "手套"]):
        return "《建筑施工作业劳动防护用品配备及使用标准》JGJ 184-2009中关于个人劳动防护用品配备和正确使用的相关条款，及项目安全管理制度"
    return base


def _resolve_risk_tail(user_input: str, ticket_type: FineTicketType) -> str:
    if ticket_type == FineTicketType.QUALITY:
        return "该问题可能导致实体质量缺陷、返工返修及后续验收受阻。"

    text = str(user_input or "")
    if any(keyword in text for keyword in ["培训证", "培训证过期", "操作证", "特种作业证", "上岗证", "证书过期", "持证", "证件过期", "无证上岗", "特种作业"]):
        return "该行为容易导致作业人员在未经有效资格确认的情况下继续操作，存在误操作和事故扩大的现实风险。"
    if any(keyword in text for keyword in ["吊装", "起重", "信号工", "司索", "塔吊", "吊物"]):
        return "该问题容易引发吊物坠落、碰撞伤害及机械伤害事故，直接威胁现场人员和设备安全。"
    if any(keyword in text for keyword in ["高处", "登高", "安全带", "临边", "洞口", "坠落"]):
        return "该问题容易导致高处坠落和物体打击事故，现场风险较高。"
    if any(keyword in text for keyword in ["临电", "配电箱", "漏保", "电缆", "接地", "电线"]):
        return "该问题容易引发触电、短路和电气火灾事故，对现场临时用电安全影响较大。"
    return "该问题继续发展容易引发人员伤害、设备损坏或现场次生事故。"


def _ensure_clause_paragraph(paragraphs: list[str], ticket_type: FineTicketType) -> list[str]:
    if not paragraphs:
        return []

    normalized = [item.strip() for item in paragraphs if item and item.strip()][:2]
    if not normalized:
        return []

    rule_reference = _default_rule_reference(ticket_type)
    if len(normalized) == 1:
        normalized.append(
            f"二、违反条款及性质：上述行为违反{rule_reference}，属于应立即整改的违规行为。"
        )
        return normalized

    second = normalized[1]
    if rule_reference not in second:
        second = second.rstrip("。") + f"，违反{rule_reference}。"
    normalized[1] = second
    return normalized


def _normalize_ai_generated_text(content: str, discovery_date: str | None, ticket_type: FineTicketType) -> str:
    normalized = str(content or "").strip()
    fact_date = _format_fact_date(discovery_date)
    rule_reference = _resolve_rule_reference(normalized, ticket_type)

    normalized = re.sub(r"\d{4}年X月X日", fact_date, normalized)
    normalized = re.sub(r"\d{4}年\d{1,2}月X日", fact_date, normalized)
    normalized = re.sub(r"\d{4}年X月\d{1,2}日", fact_date, normalized)
    normalized = re.sub(r"X年X月X日", fact_date, normalized)

    lines = [line.strip() for line in normalized.replace("\r", "").split("\n") if line.strip()]
    if lines and lines[0].startswith("一、") and fact_date not in lines[0]:
        lines[0] = lines[0].replace("一、违章事实及经过：", f"一、违章事实及经过：{fact_date}，", 1)
    if len(lines) > 1 and rule_reference not in lines[1]:
        lines[1] = lines[1].rstrip("。") + f"，违反{rule_reference}。"
    return "\n".join(lines) if lines else normalized


def _build_controlled_description(
    description: str,
    discovery_date: str | None,
    ticket_type: FineTicketType,
    rule_reference: str,
) -> list[str]:
    """Keep facts editable while keeping the legal basis and ending deterministic."""
    fact_date = _format_fact_date(discovery_date)
    ticket_name = _ticket_name(ticket_type)
    raw_lines = [
        line.strip()
        for line in str(description or "").replace("\r", "\n").split("\n")
        if line and line.strip()
    ]

    fact_lines: list[str] = []
    for line in raw_lines:
        if line.startswith("二、") or "违反条款及性质" in line:
            break
        cleaned = re.sub(r"^一、\s*违章事实及经过\s*[:：]?", "", line).strip()
        if cleaned:
            fact_lines.append(cleaned)

    fact = " ".join(fact_lines).strip()
    if not fact:
        fact = f"经现场{ticket_name}检查，发现存在违规作业行为，已责令立即整改。"
    if fact_date not in fact:
        fact = f"{fact_date}，{fact}"

    first = f"一、违章事实及经过：{fact.rstrip('。')}。"
    second = (
        f"二、违反条款及性质：上述行为不符合{rule_reference}及项目现场{ticket_name}管理要求，"
        "属于应当立即整改的违规行为。"
    )
    return [_compact_text(first, 180), _compact_text(second, 260)]


def _build_description_paragraphs(
    description: str,
    discovery_date: str | None,
    ticket_type: FineTicketType,
) -> list[str]:
    ticket_name = _ticket_name(ticket_type)
    fact_date = _format_fact_date(discovery_date)
    rule_reference = _resolve_rule_reference(description, ticket_type)

    raw_lines = [
        line.strip()
        for line in str(description or "").replace("\r", "\n").split("\n")
        if line and line.strip()
    ]

    first = ""
    second = ""
    for line in raw_lines:
        if not first and (line.startswith("一、") or "违章事实及经过" in line):
            first = line
            continue
        if not second and (line.startswith("二、") or "违反条款及性质" in line):
            second = line

    if not first:
        summary = " ".join(raw_lines).strip() or "现场发现存在违规作业行为，已责令立即整改。"
        first = f"一、违章事实及经过：{summary}"
    if fact_date not in first:
        if "一、违章事实及经过：" in first:
            first = first.replace("一、违章事实及经过：", f"一、违章事实及经过：{fact_date}，", 1)
        else:
            first = f"一、违章事实及经过：{fact_date}，{first}"

    if not second:
        second = (
            f"二、违反条款及性质：上述行为违反{rule_reference}及项目现场{ticket_name}管理要求，"
            "属于必须立即整改的违规行为，现场风险较高。"
        )
    elif rule_reference not in second:
        second = second.rstrip("。") + f"，违反{rule_reference}。"

    return [
        _compact_text(first, 122),
        _compact_text(second, 176),
    ]


def _description_too_simple(
    paragraphs: list[str],
    *,
    project_name: str = "",
    team_name: str = "",
    location: str = "",
) -> bool:
    if len(paragraphs) < 2:
        return True

    first = paragraphs[0].replace("一、违章事实及经过：", "").strip("，。；： ")
    second = paragraphs[1].replace("二、违反条款及性质：", "").strip("，。；： ")

    if len(first) < 36 or len(second) < 28:
        return True

    if first.endswith("日") or first.endswith("日，"):
        return True

    required_markers = [item for item in [project_name, team_name, location] if item]
    if required_markers and not any(marker in first for marker in required_markers):
        return True

    return False


def _compact_description(paragraphs: list[str], ticket_type: FineTicketType) -> list[str]:
    normalized = _ensure_clause_paragraph(paragraphs, ticket_type)
    compacted: list[str] = []
    for index, paragraph in enumerate(normalized[:2]):
        compacted.append(_compact_text(paragraph, 88 if index == 0 else 116))
    return compacted or normalized


def _compress_photo(source_path: Path) -> Path:
    with Image.open(source_path) as image:
        if image.mode in ("RGBA", "P"):
            image = image.convert("RGB")
        max_size = 1800
        if max(image.size) > max_size:
            ratio = max_size / max(image.size)
            image = image.resize(
                (int(image.size[0] * ratio), int(image.size[1] * ratio)),
                Image.LANCZOS,
            )
        output_path = source_path.with_name(f"{source_path.stem}_compressed.jpg")
        image.save(output_path, "JPEG", quality=86)
    return output_path


def save_uploaded_photos(files: Iterable) -> list[Path]:
    ensure_storage_dirs()
    saved_paths: list[Path] = []
    for file in files:
        if not getattr(file, "filename", ""):
            continue
        content = file.file.read()
        if len(content) > settings.MAX_UPLOAD_SIZE:
            raise ValueError(f"Upload too large. Max {settings.MAX_UPLOAD_SIZE // 1024}KB")
        _, extension = validate_image_content(content)
        output_path = PHOTOS_DIR / f"{uuid.uuid4().hex}.{extension}"
        with output_path.open("wb") as buffer:
            buffer.write(content)
        saved_paths.append(output_path)
    return saved_paths


def build_fine_document(
    *,
    number: str,
    ticket_type: FineTicketType,
    project_name: str,
    team_name: str,
    location: str,
    discovery_date: str,
    amount: Decimal,
    description: str,
    photo_paths: list[Path],
) -> tuple[str, Path]:
    ensure_storage_dirs()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run(_ticket_title(ticket_type))
    _set_run_font(title_run, size=22, bold=True, color=TITLE_RED)

    _add_bottom_line(doc)

    number_para = doc.add_paragraph()
    number_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    number_para.paragraph_format.space_before = Pt(6)
    number_para.paragraph_format.space_after = Pt(5)
    number_run = number_para.add_run(f"编号：{number}")
    _set_run_font(number_run, size=13, bold=True, color=TEXT_BLACK)

    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    meta_table.columns[0].width = Cm(4.2)
    meta_table.columns[1].width = Cm(12.6)
    meta_rows = [
        ("项目名称", project_name or "未填写"),
        ("受罚班组", team_name or "未填写"),
        ("违规部位", location or "未填写"),
        ("发现时间", _format_display_date(discovery_date)[0]),
    ]
    for row, (label, value) in zip(meta_table.rows, meta_rows):
        _set_row_height(row, 1.08)
        _set_cell_text(row.cells[0], label, size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill=LABEL_FILL)
        _set_cell_text(row.cells[1], value, size=12)
    _set_table_grid(meta_table)

    _add_section_heading(doc, "违规情况描述")
    desc_table = doc.add_table(rows=1, cols=1)
    desc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    desc_table.autofit = False
    desc_table.columns[0].width = Cm(16.8)
    _set_row_height(desc_table.rows[0], 5.7)
    desc_cell = desc_table.cell(0, 0)
    _clear_cell(desc_cell)
    desc_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _set_cell_border(
        desc_cell,
        top={"val": "single", "sz": 10, "color": BOX_BLACK},
        left={"val": "single", "sz": 10, "color": BOX_BLACK},
        bottom={"val": "single", "sz": 10, "color": BOX_BLACK},
        right={"val": "single", "sz": 10, "color": BOX_BLACK},
    )
    paragraphs = _build_description_paragraphs(description, discovery_date, ticket_type)
    for index, line in enumerate(paragraphs):
        paragraph = desc_cell.paragraphs[0] if index == 0 else desc_cell.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.16
        run = paragraph.add_run(line)
        _set_run_font(run, size=12, bold=line.startswith(("一、", "二、")))

    _add_section_heading(doc, "罚款决定")
    decision_table = doc.add_table(rows=1, cols=3)
    decision_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    decision_table.autofit = False
    decision_table.columns[0].width = Cm(4.2)
    decision_table.columns[1].width = Cm(8.2)
    decision_table.columns[2].width = Cm(4.4)
    decision_row = decision_table.rows[0]
    _set_row_height(decision_row, 1.08)
    _set_cell_text(
        decision_row.cells[0],
        "罚款金额",
        size=13,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        fill=LABEL_FILL,
    )
    amount_text = f"人民币{amount_to_chinese(amount)}（￥ {Decimal(amount):.2f} 元）"
    _set_cell_text(
        decision_row.cells[1],
        amount_text,
        size=12,
        bold=True,
        color=TITLE_RED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _set_cell_text(
        decision_row.cells[2],
        "款项须在 3 日内缴清",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _set_table_grid(decision_table)

    _add_section_heading(doc, "审批与签认")
    sign_table = doc.add_table(rows=2, cols=2)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_table.autofit = False
    sign_table.columns[0].width = Cm(4.2)
    sign_table.columns[1].width = Cm(12.6)

    submit_row = sign_table.rows[0]
    _set_row_height(submit_row, 1.08)
    _set_cell_text(
        submit_row.cells[0],
        "提报人",
        size=13,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        fill=LABEL_FILL,
    )
    _clear_cell(submit_row.cells[1])
    submit_paragraph = submit_row.cells[1].paragraphs[0]
    submit_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    submit_paragraph.paragraph_format.space_before = Pt(0)
    submit_paragraph.paragraph_format.space_after = Pt(0)

    sign_row = sign_table.rows[1]
    _set_row_height(sign_row, 1.08)
    _set_cell_text(
        sign_row.cells[0],
        "生产经理签认",
        size=13,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        fill=LABEL_FILL,
    )
    _clear_cell(sign_row.cells[1])
    sign_paragraph = sign_row.cells[1].paragraphs[0]
    sign_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sign_paragraph.paragraph_format.space_before = Pt(0)
    sign_paragraph.paragraph_format.space_after = Pt(0)
    _set_table_grid(sign_table)

    footer_date = _current_issue_date()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(8)
    footer.paragraph_format.space_after = Pt(4)
    footer_run = footer.add_run(f"开单日期：{footer_date}")
    _set_run_font(footer_run, size=12, bold=True)

    _add_bottom_line(doc)

    footnote = doc.add_paragraph()
    footnote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footnote.paragraph_format.space_before = Pt(2)
    footnote_run = footnote.add_run("本通知单一式两份，项目部留存一份，交罚班组持一份")
    _set_run_font(footnote_run, size=9)

    if photo_paths:
        doc.add_page_break()
        photos_title = doc.add_paragraph()
        photos_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        photos_title_run = photos_title.add_run("现场照片")
        _set_run_font(photos_title_run, size=18, bold=True, color=TITLE_RED)

        photo_note = doc.add_paragraph()
        photo_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_run = photo_note.add_run("以下照片作为罚款通知单附件留存。")
        _set_run_font(note_run, size=11)

        for photo_path in photo_paths:
            compressed_path = _compress_photo(photo_path)
            try:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(compressed_path), width=Cm(15.2))
            finally:
                if compressed_path.exists():
                    compressed_path.unlink()

    filename = f"{number}.docx"
    output_path = DOCS_DIR / filename
    doc.save(output_path)
    return filename, output_path


def create_rebuilt_template(template_path: Path) -> None:
    template_path.parent.mkdir(parents=True, exist_ok=True)
    filename, output_path = build_fine_document(
        number="AQ-2026-0415-001",
        ticket_type=FineTicketType.SAFETY,
        project_name="示例项目名称",
        team_name="示例受罚班组",
        location="示例违规部位",
        discovery_date="2026-04-15",
        amount=Decimal("1000"),
        description=(
            "一、违章事实及经过：现场发现作业人员未按安全要求规范作业，存在明显安全隐患，已被管理人员当场指出。\n\n"
            "二、违反条款及性质：该行为违反项目现场安全管理规定，属于必须立即整改的违规行为。"
        ),
        photo_paths=[],
    )
    if output_path != template_path:
        template_path.write_bytes(output_path.read_bytes())
