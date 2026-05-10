from copy import deepcopy
from datetime import date, datetime, time, timedelta
from pathlib import Path

import httpx
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session, joinedload

from app.api.deps import get_db
from app.api.files import get_file_user
from app.config import settings
from app.models.task import Task
from app.models.safety_log_export import SafetyLogExport
from app.models.user import User, UserRole
from app.models.work_permit import WorkPermit

router = APIRouter(prefix="/api/safety-logs", tags=["safety-logs"])

TEMPLATE_PATH = Path(__file__).resolve().parents[2] / "mobile" / "施工安全日志模板.docx"
WEEKDAY_LABELS = ["星期一", "星期二", "星期三", "星期四", "星期五", "星期六", "星期日"]
PERMIT_LABELS = {
    "hot_work_level1": "一级动火作业票",
    "hot_work_level2": "二级动火作业票",
    "hot_work_level3": "普通动火作业票",
    "height_level1": "一级高处作业票",
    "height_level2": "二级高处作业票",
    "height_level3": "三级高处作业票",
    "height_special": "特级高处作业票",
    "confined_space": "受限空间作业票",
    "lifting": "吊装作业票",
    "excavation": "动土作业票",
    "electrical": "临时用电作业票",
    "other": "其他作业票",
}
LANDSCAPE_PHOTO_SIZE_CM = (3.17, 2.38)
PORTRAIT_PHOTO_SIZE_CM = (1.59, 3.00)


def day_bounds(log_date: date) -> tuple[datetime, datetime]:
    start = datetime.combine(log_date, time.min)
    end = datetime.combine(log_date, time.max)
    return start, end


def text_value(value: object, fallback: str = "-") -> str:
    text = str(value or "").strip()
    return text or fallback


async def fetch_weather(latitude: float | None = None, longitude: float | None = None) -> str:
    location = f"{latitude},{longitude}" if latitude is not None and longitude is not None else "Chengdu"
    try:
        async with httpx.AsyncClient(timeout=5) as client:
            response = await client.get(f"https://wttr.in/{location}?format=j1&lang=zh")
            response.raise_for_status()
            data = response.json()
        current = (data.get("current_condition") or [{}])[0]
        desc = (current.get("lang_zh") or current.get("weatherDesc") or [{}])[0].get("value")
        temp = current.get("temp_C")
        return f"{desc or '多云'} {temp}℃" if temp else desc or "多云"
    except Exception:
        return "多云"


def cleanup_old_safety_logs() -> None:
    log_dir = Path(settings.UPLOAD_DIR) / "safety_logs"
    if not log_dir.exists():
        return
    cutoff = datetime.now() - timedelta(days=30)
    for path in log_dir.glob("*.docx"):
        try:
            modified_at = datetime.fromtimestamp(path.stat().st_mtime)
            if modified_at < cutoff:
                path.unlink(missing_ok=True)
        except OSError:
            continue


def upload_path(upload_url: str | None) -> Path | None:
    if not upload_url:
        return None
    relative = str(upload_url).replace("\\", "/").removeprefix("/uploads/").removeprefix("uploads/")
    path = (Path(settings.UPLOAD_DIR) / relative).resolve()
    uploads_root = Path(settings.UPLOAD_DIR).resolve()
    if uploads_root not in path.parents and path != uploads_root:
        return None
    return path if path.exists() and path.is_file() else None


def upload_urls(value: str | None) -> list[str]:
    return [url.strip() for url in str(value or "").split(",") if url.strip()]


def has_upload_in_day(value: str | None, start: datetime, end: datetime) -> bool:
    for url in upload_urls(value):
        path = upload_path(url)
        if not path:
            continue
        try:
            modified_at = datetime.fromtimestamp(path.stat().st_mtime)
        except OSError:
            continue
        if start <= modified_at <= end:
            return True
    return False


def add_cell_text(cell, text: str, *, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def clear_cell(cell) -> None:
    cell.text = ""


def delete_table(table: Table) -> None:
    table._element.getparent().remove(table._element)


def set_solid_borders(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def unlock_row_heights(table: Table) -> None:
    for row in table.rows:
        row.height = None
        tr_pr = row._tr.get_or_add_trPr()
        for height in list(tr_pr.findall(qn("w:trHeight"))):
            tr_pr.remove(height)


def append_copied_table(document: Document, table_xml) -> Table:
    document.add_page_break()
    new_table_xml = deepcopy(table_xml)
    document._body._element.append(new_table_xml)
    return Table(new_table_xml, document._body)


def add_photo(cell, upload_url: str | None, width_cm: float = 4.2, height_cm: float | None = None) -> None:
    image_path = upload_path(upload_url)
    if not image_path:
        return
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()
    try:
        if height_cm:
            run.add_picture(str(image_path), width=Cm(width_cm), height=Cm(height_cm))
        else:
            run.add_picture(str(image_path), width=Cm(width_cm))
    except Exception:
        paragraph.add_run("[照片无法插入]")


def add_photo_fit(cell, upload_url: str | None, max_width_cm: float, max_height_cm: float) -> None:
    image_path = upload_path(upload_url)
    if not image_path:
        return
    try:
        from PIL import Image

        with Image.open(image_path) as image:
            width_px, height_px = image.size
        if width_px <= 0 or height_px <= 0:
            return
        scale = min(max_width_cm / width_px, max_height_cm / height_px)
        width_cm = width_px * scale
        height_cm = height_px * scale
        add_photo(cell, upload_url, width_cm=width_cm, height_cm=height_cm)
    except Exception:
        add_photo(cell, upload_url, width_cm=max_width_cm)


def image_fit_size(upload_url: str, max_width_cm: float, max_height_cm: float) -> tuple[float, float] | None:
    image_path = upload_path(upload_url)
    if not image_path:
        return None
    try:
        from PIL import Image

        with Image.open(image_path) as image:
            width_px, height_px = image.size
        if width_px <= 0 or height_px <= 0:
            return None
        scale = min(max_width_cm / width_px, max_height_cm / height_px)
        return width_px * scale, height_px * scale
    except Exception:
        return max_width_cm, max_height_cm


def template_photo_constraint(upload_url: str) -> tuple[str, float]:
    image_path = upload_path(upload_url)
    if not image_path:
        return "width", LANDSCAPE_PHOTO_SIZE_CM[0]
    try:
        from PIL import Image

        with Image.open(image_path) as image:
            width_px, height_px = image.size
        if height_px > width_px:
            return "height", PORTRAIT_PHOTO_SIZE_CM[1]
    except Exception:
        pass
    return "width", LANDSCAPE_PHOTO_SIZE_CM[0]


def add_photos_fit(cell, value: str | None, *, max_width_cm: float, max_height_cm: float) -> None:
    urls = upload_urls(value)
    if not urls:
        return

    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    for index, url in enumerate(urls):
        image_path = upload_path(url)
        if not image_path:
            continue
        constraint, size_cm = template_photo_constraint(url)
        run = paragraph.add_run()
        try:
            if constraint == "height":
                run.add_picture(str(image_path), height=Cm(size_cm))
            else:
                run.add_picture(str(image_path), width=Cm(size_cm))
        except Exception:
            paragraph.add_run("[照片无法插入]")


def permit_value(permit: WorkPermit) -> str:
    return str(getattr(permit.type, "value", permit.type))


def permit_label(permit: WorkPermit) -> str:
    value = permit_value(permit)
    return PERMIT_LABELS.get(value, value)


def replace_text(paragraph, values: dict[str, str]) -> None:
    for run in paragraph.runs:
        text = run.text
        for key, value in values.items():
            for placeholder in (f"{{{{{key}}}}}", f"{{{key}}}", f"${{{key}}}"):
                text = text.replace(placeholder, value)
        run.text = text


def normalize_label(text: str) -> str:
    return "".join(text.replace("：", "").replace(":", "").split())


def fill_template_fields(document: Document, values: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        replace_text(paragraph, values)

    for table in document.tables:
        for row in table.rows:
            cells = row.cells
            for cell in cells:
                for paragraph in cell.paragraphs:
                    replace_text(paragraph, values)

            for index, cell in enumerate(cells[:-1]):
                label = normalize_label(cell.text)
                if label in values:
                    add_cell_text(cells[index + 1], values[label])


def risk_text(index: int, task: Task, item) -> str:
    return "\n".join(
        [
            f"隐患 {index}",
            f"区域：{task.area.name if task.area else '-'}",
            f"任务：{task.title}",
            f"风险描述：{text_value(item.risk_description)}",
            f"排查要点：{text_value(item.inspection_points)}",
            f"排查时间：{item.checked_at.strftime('%Y-%m-%d %H:%M') if item.checked_at else '-'}",
        ]
    )


def add_risk_to_cell(cell, *, index: int, task: Task, item, photo_width_cm: float) -> None:
    add_cell_text(cell, risk_text(index, task, item), size=9)
    add_photo(cell, item.photo_url, width_cm=photo_width_cm)


def fill_continuation_table(table: Table, page_items) -> None:
    set_solid_borders(table)
    unlock_row_heights(table)
    for row_index, row in enumerate(table.rows):
        text_cell, photo_cell = row.cells
        clear_cell(text_cell)
        clear_cell(photo_cell)
        if row_index < len(page_items):
            index, task, item = page_items[row_index]
            add_cell_text(text_cell, risk_text(index, task, item), size=9)
            add_photo_fit(photo_cell, item.photo_url, max_width_cm=7.2, max_height_cm=3.0)


def fill_template_table(table: Table, *, values: dict[str, str], permits, page_items) -> None:
    if len(table.rows) < 9:
        fill_template_fields(table._parent, values)
        return

    set_solid_borders(table)
    unlock_row_heights(table)
    add_cell_text(table.rows[0].cells[1], values["施工单位"])
    add_cell_text(table.rows[0].cells[2], "区域", bold=True)
    add_cell_text(table.rows[0].cells[-1], values["项目名称"])
    add_cell_text(table.rows[1].cells[1], values["日期"])
    add_cell_text(table.rows[1].cells[2], "星期", bold=True)
    add_cell_text(table.rows[1].cells[-1], values["星期"])
    add_cell_text(table.rows[2].cells[1], values["天气"])
    add_cell_text(table.rows[2].cells[2], "安全员", bold=True)
    add_cell_text(table.rows[2].cells[-1], values["安全员"])

    permit_cell = table.rows[4].cells[0]
    permit_photo_cell = table.rows[4].cells[-1]
    clear_cell(permit_cell)
    clear_cell(permit_photo_cell)
    if permits:
        for index, permit in enumerate(permits, start=1):
            paragraph = permit_cell.paragraphs[0] if index == 1 and permit_cell.paragraphs else permit_cell.add_paragraph()
            run = paragraph.add_run(
                f"{index}. {permit_label(permit)}  区域：{permit.area.name if permit.area else '-'}  负责人：{permit.responsible_person}"
            )
            run.font.size = Pt(9)
            add_photos_fit(permit_photo_cell, permit.photo_url, max_width_cm=7.4, max_height_cm=3.0)

    risk_row_template = deepcopy(table.rows[6]._tr if len(table.rows) > 6 else table.rows[-1]._tr)
    target_row_count = 6 + max(len(page_items), 1)
    while len(table.rows) < target_row_count:
        table._tbl.append(deepcopy(risk_row_template))
    while len(table.rows) > target_row_count:
        table._tbl.remove(table.rows[-1]._tr)

    for row in table.rows[6:]:
        text_cell, photo_cell = row.cells[0], row.cells[-1]
        clear_cell(text_cell)
        clear_cell(photo_cell)

    if not page_items:
        add_cell_text(table.rows[6].cells[0], "当日暂无隐患排查记录。")
        return

    for row, (index, task, item) in zip(table.rows[6:], page_items):
        text_cell, photo_cell = row.cells[0], row.cells[-1]
        add_cell_text(text_cell, risk_text(index, task, item), size=9)
        add_photos_fit(photo_cell, item.photo_url, max_width_cm=7.4, max_height_cm=3.0)


def unique_row_cells(row):
    seen = set()
    result = []
    for cell in row.cells:
        marker = id(cell._tc)
        if marker in seen:
            continue
        seen.add(marker)
        result.append(cell)
    return result


def clear_row(row) -> None:
    for cell in unique_row_cells(row):
        clear_cell(cell)


def append_row_from_template(table: Table, row_xml):
    table._tbl.append(deepcopy(row_xml))
    return table.rows[-1]


def grouped_log_sections(tasks, page_items, permits):
    sections = []
    by_task_id: dict[int, dict] = {}

    def ensure_section(task: Task | None, fallback_title: str = "") -> dict:
        task_id = getattr(task, "id", None)
        key = task_id if task_id is not None else f"standalone-{len(sections)}"
        if key in by_task_id:
            return by_task_id[key]
        title = text_value(getattr(task, "title", None), fallback_title or "当日作业")
        section = {"task": task, "title": title, "items": [], "permits": []}
        by_task_id[key] = section
        sections.append(section)
        return section

    for task in tasks or []:
        ensure_section(task)

    for task, item in page_items or []:
        ensure_section(task)["items"].append((task, item))

    for permit in permits or []:
        permit_task = getattr(permit, "task", None)
        if permit_task is not None:
            ensure_section(permit_task)["permits"].append(permit)
            continue

        matching = next(
            (
                section
                for section in sections
                if section["task"] is not None
                and getattr(section["task"], "area_id", None) == getattr(permit, "area_id", None)
            ),
            None,
        )
        if matching is None:
            matching = ensure_section(None, getattr(permit, "description", "") or permit_label(permit))
        matching["permits"].append(permit)

    return [section for section in sections if section["items"] or section["permits"]]


def value_at(values: dict[str, str], index: int, fallback: str = "-") -> str:
    try:
        return list(values.values())[index]
    except IndexError:
        return fallback


def fill_template_table(table: Table, *, values: dict[str, str], tasks, permits, page_items) -> None:
    if len(table.rows) < 7:
        fill_template_fields(table._parent, values)
        return

    set_solid_borders(table)
    unlock_row_heights(table)
    add_cell_text(table.rows[0].cells[1], value_at(values, 0))
    add_cell_text(table.rows[0].cells[2], "区域", bold=True)
    add_cell_text(table.rows[0].cells[-1], value_at(values, 1))
    add_cell_text(table.rows[1].cells[1], value_at(values, 3))
    add_cell_text(table.rows[1].cells[2], "星期", bold=True)
    add_cell_text(table.rows[1].cells[-1], value_at(values, 4))
    add_cell_text(table.rows[2].cells[1], value_at(values, 5))
    add_cell_text(table.rows[2].cells[2], "安全员", bold=True)
    add_cell_text(table.rows[2].cells[-1], value_at(values, 6))

    title_row_template = deepcopy(table.rows[3]._tr)
    permit_header_template = deepcopy(table.rows[3]._tr)
    permit_content_template = deepcopy(table.rows[4]._tr)
    risk_header_template = deepcopy(table.rows[5]._tr)
    risk_row_template = deepcopy(table.rows[6]._tr)

    while len(table.rows) > 3:
        table._tbl.remove(table.rows[-1]._tr)

    sections = grouped_log_sections(tasks, page_items, permits)
    if not sections:
        sections = [{"task": None, "title": "当日作业", "items": [], "permits": []}]

    for section in sections:
        title_row = append_row_from_template(table, title_row_template)
        clear_row(title_row)
        add_cell_text(title_row.cells[0], section["title"], bold=True, size=10)

        permit_header_row = append_row_from_template(table, permit_header_template)
        clear_row(permit_header_row)
        add_cell_text(permit_header_row.cells[0], "作业票据", bold=True, size=10)

        permit_row = append_row_from_template(table, permit_content_template)
        clear_row(permit_row)
        permit_cell = permit_row.cells[0]
        permit_photo_cell = permit_row.cells[-1]
        for index, permit in enumerate(section["permits"], start=1):
            paragraph = permit_cell.paragraphs[0] if index == 1 and permit_cell.paragraphs else permit_cell.add_paragraph()
            run = paragraph.add_run(
                f"{index}. {permit_label(permit)}  区域：{permit.area.name if permit.area else '-'}  责任人：{permit.responsible_person}"
            )
            run.font.size = Pt(9)
            add_photos_fit(permit_photo_cell, permit.photo_url, max_width_cm=7.4, max_height_cm=3.0)

        risk_header_row = append_row_from_template(table, risk_header_template)
        clear_row(risk_header_row)
        add_cell_text(risk_header_row.cells[0], "隐患排查", bold=True, size=10)
        add_cell_text(risk_header_row.cells[-1], "整改照片", bold=True, size=10)

        if not section["items"]:
            risk_row = append_row_from_template(table, risk_row_template)
            clear_row(risk_row)
            add_cell_text(risk_row.cells[0], "当日暂无隐患排查记录。", size=9)
            continue

        for local_index, (task, item) in enumerate(section["items"], start=1):
            risk_row = append_row_from_template(table, risk_row_template)
            clear_row(risk_row)
            add_cell_text(risk_row.cells[0], risk_text(local_index, task, item), size=9)
            add_photos_fit(risk_row.cells[-1], item.photo_url, max_width_cm=7.4, max_height_cm=3.0)


def user_display_name(user: User | None) -> str:
    if not user:
        return "-"
    return user.real_name or user.username or "-"


def can_generate_for(current_user: User, target_user: User) -> bool:
    return current_user.id == target_user.id or current_user.role == UserRole.ADMIN


def resolve_log_subject(db: Session, current_user: User, target_user_id: int | None) -> User:
    if target_user_id is None or target_user_id == current_user.id:
        return current_user
    if current_user.role != UserRole.ADMIN:
        raise HTTPException(status_code=403, detail="Only admins can generate logs for other users")
    target_user = db.query(User).filter(User.id == target_user_id).first()
    if not target_user:
        raise HTTPException(status_code=404, detail="Safety inspector not found")
    return target_user


def collect_log_data(db: Session, current_user: User, log_date: date):
    start, end = day_bounds(log_date)
    candidate_tasks = (
        db.query(Task)
        .options(joinedload(Task.area), joinedload(Task.assignee), joinedload(Task.checklist_items))
        .filter(Task.assignee_id == current_user.id)
        .filter(Task.created_at <= end)
        .order_by(Task.created_at.desc())
        .all()
    )

    items = []
    included_tasks: dict[int, Task] = {}
    for task in candidate_tasks:
        for item in task.checklist_items or []:
            in_checked_day = item.checked_at and start <= item.checked_at <= end
            in_task_day = task.created_at and start <= task.created_at <= end
            in_photo_day = has_upload_in_day(item.photo_url, start, end)
            if in_checked_day or in_task_day or in_photo_day:
                items.append((task, item))
                included_tasks[task.id] = task

    candidate_permits = (
        db.query(WorkPermit)
        .options(joinedload(WorkPermit.area), joinedload(WorkPermit.applicant), joinedload(WorkPermit.task))
        .filter(WorkPermit.applicant_id == current_user.id)
        .filter(WorkPermit.photo_url.isnot(None))
        .order_by(WorkPermit.created_at.desc())
        .all()
    )
    permits = [
        permit
        for permit in candidate_permits
        if (permit.created_at and start <= permit.created_at <= end)
        or has_upload_in_day(permit.photo_url, start, end)
    ]
    for permit in permits:
        if permit.task_id:
            task = next((task for task in candidate_tasks if task.id == permit.task_id), None)
            if task:
                included_tasks[task.id] = task

    return list(included_tasks.values()), items, permits


def build_docx(
    *,
    output_path: Path,
    log_date: date,
    weather: str,
    current_user: User,
    tasks,
    items,
    permits,
) -> None:
    inspector_name = current_user.real_name or current_user.username
    area_names = sorted({task.area.name for task in tasks if task.area and task.area.name})
    project_name = "、".join(area_names) if area_names else "风险区域"
    info_values = {
        "施工单位": "四川华庭",
        "项目名称": project_name,
        "巡查日期": log_date.strftime("%Y年%m月%d日"),
        "日期": log_date.strftime("%Y年%m月%d日"),
        "星期": WEEKDAY_LABELS[log_date.weekday()],
        "天气": weather,
        "安全员": inspector_name,
        "任务执行人": inspector_name,
        "作业票据": "、".join(permit_label(permit) for permit in permits) or "无",
        "隐患数量": str(len(items)),
    }

    using_template = TEMPLATE_PATH.exists()
    document = Document(str(TEMPLATE_PATH)) if using_template else Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)

    if using_template:
        if document.tables:
            fill_template_table(
                document.tables[0],
                values=info_values,
                tasks=tasks,
                permits=permits,
                page_items=items,
            )
        if len(document.tables) > 1:
            for table in list(document.tables[1:]):
                delete_table(table)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        return
    else:
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("施工安全日志")
        run.bold = True
        run.font.size = Pt(18)

        info = document.add_table(rows=4, cols=4)
        info.style = "Table Grid"
        info_data = [
            ("施工单位", info_values["施工单位"], "项目名称", info_values["项目名称"]),
            ("巡查日期", info_values["巡查日期"], "星期", info_values["星期"]),
            ("天气", info_values["天气"], "安全员", info_values["安全员"]),
            ("作业票据", info_values["作业票据"], "隐患数量", info_values["隐患数量"]),
        ]
        for row, row_values in zip(info.rows, info_data):
            for cell, value in zip(row.cells, row_values):
                add_cell_text(cell, value, bold=value in {"施工单位", "项目名称", "巡查日期", "星期", "天气", "安全员", "作业票据", "隐患数量"})

    if permits:
        document.add_paragraph("")
        heading = document.add_paragraph()
        heading.add_run("作业许可照片").bold = True
        permit_table = document.add_table(rows=0, cols=2)
        permit_table.style = "Table Grid"
        for permit in permits:
            row = permit_table.add_row()
            add_cell_text(
                row.cells[0],
                f"{permit_label(permit)}\n区域：{permit.area.name if permit.area else '-'}\n责任人：{permit.responsible_person}",
                bold=False,
            )
            add_photo(row.cells[1], permit.photo_url)

    document.add_paragraph("")
    hidden_heading = document.add_paragraph()
    hidden_heading.add_run("隐患排查记录").bold = True

    if not items:
        document.add_paragraph("当日暂无隐患排查记录。")
    else:
        for index, (task, item) in enumerate(items, start=1):
            if index > 1 and (index - 1) % 4 == 0:
                document.add_page_break()
            table = document.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            cell = table.rows[0].cells[0]
            add_cell_text(
                cell,
                "\n".join(
                    [
                        f"隐患 {index}",
                        f"区域：{task.area.name if task.area else '-'}",
                        f"任务：{task.title}",
                        f"风险描述：{text_value(item.risk_description)}",
                        f"排查要点：{text_value(item.inspection_points)}",
                        f"排查时间：{item.checked_at.strftime('%Y-%m-%d %H:%M') if item.checked_at else '-'}",
                    ]
                ),
                size=10,
            )
            add_photo(cell, item.photo_url, width_cm=9.0)
            document.add_paragraph("")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


@router.get("/generate")
async def generate_safety_log(
    log_date: date | None = Query(default=None),
    lat: float | None = Query(default=None),
    lon: float | None = Query(default=None),
    user_id: int | None = Query(default=None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_file_user),
):
    log_date = log_date or date.today()
    cleanup_old_safety_logs()
    subject_user = resolve_log_subject(db, current_user, user_id)
    tasks, items, permits = collect_log_data(db, subject_user, log_date)
    weather = await fetch_weather(lat, lon)
    filename = f"safety-log-{subject_user.id}-{log_date.isoformat()}.docx"
    output_path = Path(settings.UPLOAD_DIR) / "safety_logs" / filename
    build_docx(
        output_path=output_path,
        log_date=log_date,
        weather=weather,
        current_user=subject_user,
        tasks=tasks,
        items=items,
        permits=permits,
    )
    relative_file_path = f"/uploads/safety_logs/{filename}"
    db.add(
        SafetyLogExport(
            subject_user_id=subject_user.id,
            exported_by_id=current_user.id,
            log_date=log_date,
            file_path=relative_file_path,
        )
    )
    db.commit()
    return FileResponse(
        output_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"施工安全日志-{user_display_name(subject_user)}-{log_date.isoformat()}.docx",
    )


@router.get("/history")
def list_safety_log_history(
    limit: int = Query(default=50, ge=1, le=200),
    user_id: int | None = Query(default=None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_file_user),
):
    query = (
        db.query(SafetyLogExport)
        .options(joinedload(SafetyLogExport.subject_user), joinedload(SafetyLogExport.exported_by))
        .order_by(SafetyLogExport.created_at.desc())
    )
    if current_user.role != UserRole.ADMIN:
        query = query.filter(SafetyLogExport.subject_user_id == current_user.id)
    elif user_id:
        query = query.filter(SafetyLogExport.subject_user_id == user_id)

    rows = query.limit(limit).all()
    return [
        {
            "id": row.id,
            "log_date": row.log_date.isoformat() if row.log_date else "",
            "file_path": row.file_path,
            "created_at": row.created_at.isoformat() if row.created_at else "",
            "subject_user": {
                "id": row.subject_user.id if row.subject_user else None,
                "username": row.subject_user.username if row.subject_user else "",
                "real_name": user_display_name(row.subject_user),
            },
            "exported_by": {
                "id": row.exported_by.id if row.exported_by else None,
                "username": row.exported_by.username if row.exported_by else "",
                "real_name": user_display_name(row.exported_by),
            },
        }
        for row in rows
    ]
